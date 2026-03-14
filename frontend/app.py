import streamlit as st
import requests
import pandas as pd
import pickle
import datetime
import re
import os
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# --- 1. PAGE CONFIG & SESSION STATE ---
st.set_page_config(
    page_title="Legal Assistant Pro & Audit System", 
    page_icon="⚖️", 
    layout="wide",
    initial_sidebar_state="expanded"
)
API = "http://localhost:8000"

# Initialize Session States (Ensuring full coverage of project data)
if 'testimonies' not in st.session_state:
    st.session_state.testimonies = []
if 'prediction_data' not in st.session_state:
    st.session_state.prediction_data = None
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'current_case_data' not in st.session_state:
    st.session_state.current_case_data = None
if 'case_id' not in st.session_state:
    st.session_state.case_id = None
if 'current_filename' not in st.session_state:
    st.session_state.current_filename = "No Case Loaded"

# --- 2. LOAD RESOURCES (Local ML Model) ---
@st.cache_resource
def load_resources():
    """Loads the local Scikit-Learn model and the historical justice dataset."""
    try:
        model = pickle.load(open("simple_facts_model.pkl", "rb"))
        df = pd.read_csv("justice.csv")
        df['facts'] = df['facts'].fillna('')
        return model, df
    except Exception as e:
        st.sidebar.error(f"Error loading local ML resources: {e}")
        return None, None

model_local, df_historical = load_resources()

# --- 3. HELPER FUNCTIONS ---
def extract_points(text):
    """Uses regex and keyword matching to find critical witness observations."""
    sentences = re.split(r'(?<=[.!?]) +', text)
    keywords = ["saw", "heard", "witnessed", "observed", "presence", "identity", "stated", "occurred", "incident"]
    extracted = [s.strip() for s in sentences if any(w in s.lower() for w in keywords) and len(s) > 15]
    return extracted

def find_reference(user_input):
    """Calculates TF-IDF similarity to find the most relevant historical case."""
    if df_historical is None:
        return "N/A", 0, 0.0
    tfidf = TfidfVectorizer(stop_words='english')
    all_facts = df_historical['facts'].tolist() + [user_input]
    tfidf_matrix = tfidf.fit_transform(all_facts)
    similarities = cosine_similarity(tfidf_matrix[-1], tfidf_matrix[:-1])
    idx = similarities.argmax()
    return df_historical.iloc[idx]['ID'], df_historical.iloc[idx]['first_party_winner'], similarities[0][idx]

def create_word_report(data, testimonies):
    """Generates a professional .docx report containing all AI and ML findings."""
    doc = Document()
    
    # Title Section
    title = doc.add_heading('⚖️ Comprehensive Case Audit & Legal Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Report Generated: {datetime.date.today()}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("_" * 50)

    # Section 1: ML Prediction
    doc.add_heading('1. Machine Learning Prediction (Local Model)', level=1)
    p1 = doc.add_paragraph()
    p1.add_run(f"Predicted Outcome: ").bold = True
    p1.add_run(f"{data.get('pred_text', 'N/A')}")
    
    p2 = doc.add_paragraph()
    p2.add_run(f"Confidence Level: ").bold = True
    p2.add_run(f"{data.get('prob', 0.0):.2f}%")
    
    p3 = doc.add_paragraph()
    p3.add_run(f"Similar Historical Reference: ").bold = True
    p3.add_run(f"{data.get('ref_id', 'N/A')} (Outcome: {data.get('ref_out', 'N/A')})")

    # Section 2: Timeline
    doc.add_heading('2. Timeline & Procedural Lapse', level=1)
    doc.add_paragraph(f"Date of Incident: {data.get('inc_date', 'N/A')}")
    doc.add_paragraph(f"Scheduled Hearing: {data.get('hear_date', 'N/A')}")
    
    lapse_para = doc.add_paragraph()
    lapse_para.add_run(f"Total Lapse (Delay): {data.get('lapse', 0)} days").bold = True

    # Section 3: Compliance
    doc.add_heading('3. Investigation & Legal Compliance', level=1)
    doc.add_paragraph(f"Police Investigation Verified: {'[X] COMPLETE' if data.get('p_ready') else '[ ] PENDING'}")
    doc.add_paragraph(f"Lawyer Documentation Verified: {'[X] COMPLETE' if data.get('l_ready') else '[ ] PENDING'}")

    # Section 4: Testimonies
    doc.add_heading('4. Witness Testimonies & Evidence Analysis', level=1)
    if not testimonies:
        doc.add_paragraph("No witness statements were provided for this audit.")
    else:
        for t in testimonies:
            doc.add_heading(f"Witness Name: {t['name']}", level=2)
            doc.add_heading("Full Statement:", level=3)
            doc.add_paragraph(t['full_text'])
            doc.add_heading("Extracted Key Evidence Points:", level=3)
            for pt in t['points']:
                doc.add_paragraph(pt, style='List Bullet')

    # Save to buffer
    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()

# --- 4. SIDEBAR: LEGAL VAULT (SEMANTIC SEARCH) ---
st.sidebar.title("📚 Legal Vault")
st.sidebar.markdown(f"**Status:** Case Loaded" if st.session_state.case_id else "**Status:** No Case Active")
st.sidebar.info(f"📁 Active Case: {st.session_state.current_filename}")

st.sidebar.divider()
search_q = st.sidebar.text_input("🔍 Semantic Vault Search", placeholder="e.g. child marriage guidelines")

if search_q:
    try:
        res = requests.get(f"{API}/search", params={"q": search_q}).json()
        results = res.get("results", [])
        if results:
            st.sidebar.write(f"**Found {len(results)} relevant cases:**")
            for r in results:
                # Clickable result to load case into session
                if st.sidebar.button(f"📄 {r['filename']}", key=f"side_{r['case_id']}"):
                    st.session_state.case_id = r['case_id']
                    st.session_state.current_filename = r['filename']
                    st.session_state.current_case_data = {"report": {"summary": r['summary_preview']}}
                    st.session_state.messages = [] 
                    st.rerun()
                st.sidebar.caption(f"Context Similarity: {int(r['score']*100)}%")
        else:
            st.sidebar.warning("No matching scenarios found in your vault.")
    except Exception as e:
        st.sidebar.error("Backend Search Offline")

if st.sidebar.button("🗑️ Reset All Session Data", use_container_width=True):
    st.session_state.testimonies = []
    st.session_state.prediction_data = None
    st.session_state.messages = []
    st.session_state.case_id = None
    st.session_state.current_filename = "No Case Loaded"
    st.rerun()

# --- 5. MAIN UI TABS ---
st.title("⚖️ Supreme Court Management & AI Assistant")
st.markdown("#### B.Tech Data Science Project - VR Siddhartha Engineering College")

tabs = st.tabs(["📄 AI Case Analysis", "🔍 ML Prediction", "💬 Legal Chatbot", "👥 Witness Analysis", "📄 Export Report"])

# --- TAB 1: AI CASE ANALYSIS ---
with tabs[0]:
    st.subheader("AI Case Extraction & Summary")
    up = st.file_uploader("Upload Case PDF", type="pdf", key="main_pdf")
    
    if up and st.button("Analyze & Save to Cloud Vault"):
        with st.spinner("Gemini AI is analyzing and checking for duplicates..."):
            try:
                r = requests.post(f"{API}/analyze", files={"file": up})
                data = r.json()
                
                # Handling Repeated Case response
                if data.get("status") == "repeated":
                    st.warning(f"⚠️ {data['explanation']}")
                    st.session_state.case_id = data['case_id']
                    st.session_state.current_filename = data['filename']
                    st.session_state['current_case_data'] = data
                elif "case_id" in data:
                    st.session_state.case_id = data['case_id']
                    st.session_state.current_filename = up.name
                    st.session_state['current_case_data'] = data
                    st.success(f"Case Analyzed Successfully! Cloud ID: {data['case_id']}")
                
                if st.session_state.current_case_data:
                    st.divider()
                    st.subheader("📝 Case Summary")
                    st.write(st.session_state.current_case_data['report']['summary'])
                    st.info(f"Predicted Backend Outcome: {data.get('prediction', 'Not Calculated')}")
            except Exception as e:
                st.error(f"Backend Error: {e}")

# --- TAB 2: ML PREDICTION (MANUAL FACTS) ---
with tabs[1]:
    st.subheader("Legal Case Prediction & Similarity (Historical Data)")
    case_facts_input = st.text_area("Input Case Facts for ML Analysis:", height=200, placeholder="Paste case facts here to run local ML prediction...")
    
    if st.button("Calculate Outcome Probability", type="primary"):
        if case_facts_input:
            if model_local is None:
                st.error("ML Model files (.pkl/.csv) not found in directory.")
            else:
                pred = model_local.predict([case_facts_input])[0]
                prob = model_local.predict_proba([case_facts_input])[0][1] * 100
                ref_id, ref_winner, sim_score = find_reference(case_facts_input)
                
                st.session_state.prediction_data = {
                    "pred_text": "First Party Wins" if pred == 1 else "First Party Loses",
                    "prob": prob,
                    "ref_id": ref_id,
                    "ref_out": "Win" if ref_winner == 1 else "Loss"
                }
                
                st.divider()
                if pred == 1:
                    st.success(f"### ML Prediction: Win ({prob:.2f}%)")
                else:
                    st.error(f"### ML Prediction: Loss ({prob:.2f}%)")
                st.info(f"**Most Similar Historical Case:** {ref_id} | Similarity Score: {sim_score:.4f}")

# --- TAB 3: LEGAL CHATBOT ---
with tabs[2]:
    st.header(f"💬 Legal Consultant: {st.session_state.current_filename}")
    if not st.session_state.case_id:
        st.warning("Please load a case in Tab 1 or from the Sidebar Vault to begin.")
    else:
        # Chat Interface
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        if chat_prompt := st.chat_input("Ask about evidence, laws, or case details..."):
            st.chat_message("user").markdown(chat_prompt)
            st.session_state.messages.append({"role": "user", "content": chat_prompt})

            with st.spinner("AI Lawyer is reviewing case history..."):
                payload = {"case_id": st.session_state.case_id, "query": chat_prompt}
                try:
                    res = requests.post(f"{API}/chat_with_history", json=payload).json()
                    answer = res.get('answer', "System could not generate a response.")
                    with st.chat_message("assistant"):
                        st.markdown(answer)
                    st.session_state.messages.append({"role": "assistant", "content": answer})
                except Exception as e:
                    st.error(f"Chat Backend Error: {e}")

# --- TAB 4: WITNESS ANALYSIS ---
with tabs[3]:
    st.subheader("Witness Testimony & Evidence Management")
    w_col1, w_col2 = st.columns([1, 2])
    with w_col1:
        witness_name = st.text_input("Witness Full Name:")
    with w_col2:
        witness_statement = st.text_area("Full Testimony Statement:", height=150)
    
    if st.button("Add Statement & Extract Points"):
        if witness_name and witness_statement:
            evidence_points = extract_points(witness_statement)
            st.session_state.testimonies.append({
                "name": witness_name, 
                "points": evidence_points,
                "full_text": witness_statement
            })
            st.success(f"Evidence extracted from {witness_name}.")
            st.rerun()
    
    st.divider()
    st.markdown("### 📋 Recorded Statements")
    for idx, test in enumerate(st.session_state.testimonies):
        with st.expander(f"Witness: {test['name']}"):
            st.write("**Statement:**")
            st.write(test['full_text'])
            st.write("**Key Extracted Evidence:**")
            for point in test['points']:
                st.write(f"- {point}")
            if st.button("Remove Entry", key=f"del_{idx}"):
                st.session_state.testimonies.pop(idx)
                st.rerun()

# --- TAB 5: EXPORT REPORT ---
with tabs[4]:
    st.subheader("Justice System Audit & Final Export")
    col_r1, col_r2 = st.columns(2)
    with col_r1:
        p_check = st.checkbox("Investigation Report Uploaded")
        l_check = st.checkbox("Counsel Documentation Filed")
    with col_r2:
        date_inc = st.date_input("Incident Date", datetime.date(2023, 1, 1))
        date_hear = st.date_input("Court Hearing Date")
    
    days_lapse = (date_hear - date_inc).days
    st.metric("Total Delay (Days)", f"{days_lapse} Days", delta="Justice Delayed" if days_lapse > 30 else "Normal")
    
    if st.button("Generate Final Audit Word Document", type="primary"):
        if st.session_state.prediction_data:
            final_report_data = st.session_state.prediction_data
            final_report_data.update({
                "inc_date": date_inc, 
                "hear_date": date_hear,
                "lapse": days_lapse, 
                "p_ready": p_check, 
                "l_ready": l_check
            })
            
            report_bytes = create_word_report(final_report_data, st.session_state.testimonies)
            st.download_button(
                label="📥 Download Legal Audit Report",
                data=report_bytes,
                file_name=f"Audit_Report_{st.session_state.current_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("Please run a case prediction in Tab 2 before generating the report.")